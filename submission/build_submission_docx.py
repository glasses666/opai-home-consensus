from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "元界视创-欧派AI家装共识工作台-参赛方案.md"
OUTPUT = ROOT / "元界视创-欧派AI家装共识工作台-参赛方案.docx"
EVIDENCE = ROOT / "evidence" / "raw"

BURGUNDY = "8F3D57"
TAN = "D8B47C"
INK = "272321"
MUTED = "716965"
PAPER = "FAF7F2"
PALE = "F2E8E4"
FONT_NAME = "Arial Unicode MS"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_hyperlink(paragraph, text, url, color=BURGUNDY):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color)
    props.append(clr)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\(https?://[^)]+\)|`[^`]+`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(BURGUNDY)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_picture(doc, filename, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(EVIDENCE / filename), width=Inches(width))
    add_caption(doc, caption)


def add_two_pictures(doc, left, right, captions):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, filename in zip(table.rows[0].cells, (left, right)):
        set_cell_margins(cell, 40, 40, 40, 40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(EVIDENCE / filename), width=Inches(3.0))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(captions)
    run.italic = True
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 14),
        ("Heading 1", 20, BURGUNDY, 18, 8),
        ("Heading 2", 15, BURGUNDY, 14, 6),
        ("Heading 3", 12, INK, 11, 4),
        ("Heading 4", 10.5, BURGUNDY, 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "元界视创 × 欧派家居｜AI 家装共识工作台"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI 先理解 · 家庭再调整 · 共识可交接   |   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("AI 先锋未来人才大赛｜40 强赛参赛方案")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BURGUNDY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(22)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("AI 家装共识工作台")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("把“我想要”变成全家看得见、改得动、交得出的空间共识")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(BURGUNDY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.add_run().add_picture(str(EVIDENCE / "07-product-public-demo.png"), width=Inches(6.1))

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Inches(1.3)
    meta.columns[1].width = Inches(4.7)
    values = (("参赛队伍", "元界视创"), ("企业命题", "欧派家居企业命题"), ("提交日期", "2026 年 8 月 16 日"))
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], PALE)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = quote.add_run("AI 不替家庭做决定；AI 让每个决定都有来源。")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("阅读导航", level=1)
    for number, title, desc in (
        ("01", "参赛方案信息卡", "队伍、摘要、分工与飞书 AI 能力"),
        ("02", "方案成果展示", "痛点、创新、技术、场景与价值"),
        ("03", "自由展示区", "我们如何理解 AI 先锋"),
        ("04", "附录", "幕后故事、证据来源与能力边界"),
    ):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(0.8)
        table.columns[1].width = Inches(5.5)
        left, right = table.rows[0].cells
        left.text = number
        right.text = f"{title}\n{desc}"
        set_cell_shading(left, BURGUNDY)
        left.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        left.paragraphs[0].runs[0].font.size = Pt(16)
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        right.paragraphs[0].runs[0].bold = True
        set_cell_margins(left, 180, 80, 180, 80)
        set_cell_margins(right, 120, 180, 120, 180)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
    doc.add_page_break()


def parse_markdown(doc, lines):
    in_code = False
    code_lines = []
    i = 0
    inserted = set()
    pending_report = False
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_cell_shading(table.cell(0, 0), "F0ECE7")
                p = table.cell(0, 0).paragraphs[0]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Menlo"
                run.font.size = Pt(8.8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line or line == "---":
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[\s:|-]+\|$", lines[i].strip()):
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    add_inline(cell.paragraphs[0], value)
                    set_cell_margins(cell)
                    if r_idx == 0:
                        set_cell_shading(cell, BURGUNDY)
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    elif r_idx % 2 == 0:
                        set_cell_shading(cell, "F8F2ED")
            set_repeat_table_header(table.rows[0])
            doc.add_paragraph()
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).replace("**", "")
            if level == 1 and title.startswith("元界视创"):
                i += 1
                continue
            if pending_report and "可落地与可推广路径" in title:
                add_picture(doc, "06-oppein-annual-report-ai-page.png", "图 5｜欧派家居 2025 年年度报告第 18 页：智家云、设计生产一体化与 AI 试点方向", width=5.6)
                inserted.add("report")
                pending_report = False
            doc.add_heading(title, level=min(level, 4))
            if "命题场景、问题与痛点" in title and "survey" not in inserted:
                add_picture(doc, "03-consumer-survey.png", "图 1｜山东家装家居消费调查报道中的沟通与落地数据（腾讯新闻页面截图）")
                add_two_pictures(doc, "04-xhs-communication-case.png", "05-xhs-contract-case.png", "图 2｜公开社交平台中的两类定性用户声音：反复改稿疲惫与需求/合同边界摩擦")
                inserted.add("survey")
            elif "三个可演示场景" in title and "product" not in inserted:
                add_picture(doc, "07-product-public-demo.png", "图 3｜阿里云公网部署的整屋 3D 与设计助理界面（2026-08-16）")
                inserted.add("product")
            elif "AI 技术链路" in title and "aily" not in inserted:
                add_picture(doc, "02-feishu-aily-capabilities.png", "图 4｜飞书官网对 Aily 工作流能力的公开介绍")
                inserted.add("aily")
            elif "对欧派" in title and "report" not in inserted:
                pending_report = True
            i += 1
            continue

        if line.startswith(">"):
            table = doc.add_table(rows=1, cols=1)
            set_cell_shading(table.cell(0, 0), PALE)
            p = table.cell(0, 0).paragraphs[0]
            add_inline(p, line.lstrip("> "))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(BURGUNDY)
            set_cell_margins(table.cell(0, 0), 180, 220, 180, 220)
            doc.add_paragraph()
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or numbered:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            add_inline(p, (bullet or numbered).group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line.replace("  ", " "))
        i += 1


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)
    add_contents(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    parse_markdown(doc, lines)

    def style_runs(container):
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                run.font.name = FONT_NAME
                fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                    fonts.set(qn(f"w:{key}"), FONT_NAME)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    style_runs(cell)

    style_runs(doc)
    for current_section in doc.sections:
        style_runs(current_section.header)
        style_runs(current_section.footer)

    props = doc.core_properties
    props.title = "元界视创｜AI 家装共识工作台参赛方案"
    props.subject = "AI 先锋未来人才大赛 40 强赛"
    props.author = "元界视创"
    props.keywords = "欧派家居, 飞书 Aily, 家装共识, 3D, SceneCommand"
    props.comments = "不含任何 API 密钥、用户令牌或企业私密凭据。"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
